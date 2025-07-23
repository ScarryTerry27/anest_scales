import datetime
import io

import streamlit as st
from docx import Document

from backend.scales.ariscat import check_ariscat
from backend.scales.caprini import caprini_score_and_risk
from backend.scales.el_ganzouri import check_elganzouri
from backend.scales.index_lee import lee_score_and_risk
from backend.scales.obesity import calculate_bmi, calculate_idmt, calculate_tmt, calculate_cmt
from backend.scales.soba_recomendation import is_high_soba_risk, stopbang_score
from frontend.components.diagnostic import show_diagnostic
from frontend.components.report import show_characteristics_of_body_weight
from frontend.components.scales import show_an_ream_risk
from frontend.components.treatment import show_treatment, show_ivl, show_recommendations_pp


def show_button_down():
    col_prev, col_next = st.columns(2, gap="large", border=False)
    if st.session_state["stage"] and col_prev.button("⬅️ Назад", use_container_width=True):
        prev_router()


def next_router():
    st.session_state["stage"] = (st.session_state["stage"] + 1) % 7
    st.rerun()


def prev_router():
    st.session_state["stage"] = abs(st.session_state["stage"] + 7 - 1) % 7
    st.rerun()


def check_scales():
    gender = "man" if st.session_state.patient_data["Пол"] == "Муж" else "woman"

    caprini = caprini_score_and_risk(
        {k: v for k, v in st.session_state.patient_data.items() if "Caprini" in k}
    )
    elganzouri = check_elganzouri(
        {k: v for k, v in st.session_state.patient_data.items() if "ElGanzouri" in k}
    )
    lee = lee_score_and_risk(
        st.session_state.patient_data["LEE_Операция"],
        st.session_state.patient_data["LEE_ИБС"],
        st.session_state.patient_data["LEE_ХСН"],
        st.session_state.patient_data["LEE_ОНМК"],
        st.session_state.patient_data["LEE_СД"],
        st.session_state.patient_data["LEE_Креатинин"],
    )
    ariscat = check_ariscat(
        st.session_state.patient_data["Возраст"],
        st.session_state.patient_data["spo2"],
        st.session_state.patient_data["ARISCAT_Инфекция"] != "Нет (0 баллов)",
        st.session_state.patient_data["ARISCAT_Анемия"] != "Нет (0 баллов)",
        st.session_state.patient_data["ARISCAT_Разрез"],
        st.session_state.patient_data["ARISCAT_Длительность"],
        st.session_state.patient_data["ARISCAT_Экстренная"] != "Нет (0 баллов)",
    )
    bmi = calculate_bmi(st.session_state.patient_data["Рост"], st.session_state.patient_data["Вес"], )
    idmt = calculate_idmt(st.session_state.patient_data["Рост"], gender)
    tmt = calculate_tmt(st.session_state.patient_data["Вес"], bmi[0], gender)
    cmt = calculate_cmt(st.session_state.patient_data["Вес"], idmt)

    stopbang = stopbang_score(
        st.session_state.patient_data["STOPBANG_Храп"],
        st.session_state.patient_data["STOPBANG_Сонливость"],
        st.session_state.patient_data["STOPBANG_Апноэ"],
        st.session_state.patient_data["STOPBANG_Давление"],
        bmi[0],
        st.session_state.patient_data["Возраст"],
        st.session_state.patient_data["STOPBANG_Шея"],
        gender
    )
    soba = is_high_soba_risk(
        st.session_state.patient_data["SOBA_Функция"],
        st.session_state.patient_data["SOBA_ЭКГ"],
        st.session_state.patient_data["SOBA_АГ"],
        st.session_state.patient_data["spo2"],
        st.session_state.patient_data["SOBA_CO2"],
        st.session_state.patient_data["SOBA_ТГВ"],
        stopbang
    )
    st.session_state["scales"]["caprini"] = caprini
    st.session_state["scales"]["elganzouri"] = elganzouri
    st.session_state["scales"]["lee"] = lee
    st.session_state["scales"]["ariscat"] = ariscat
    st.session_state["scales"]["bmi"] = bmi
    st.session_state["scales"]["idmt"] = idmt
    st.session_state["scales"]["tmt"] = tmt
    st.session_state["scales"]["cmt"] = cmt
    st.session_state["scales"]["stopbang"] = stopbang
    st.session_state["scales"]["soba"] = soba


def show_scales():
    check_scales()

    show_characteristics_of_body_weight()

    show_an_ream_risk()

    show_diagnostic()

    show_treatment()

    show_ivl()

    show_recommendations_pp()

def load_docx():
    patient = st.session_state.patient_data
    scales = st.session_state.scales

    document = Document()
    document.add_heading('🧾 Отчёт по пациенту', level=1)

    document.add_paragraph(f'Дата: {datetime.datetime.now().strftime("%d.%m.%Y %H:%M")}')
    document.add_paragraph('')

    # Общая информация
    document.add_heading('🔹 Общие сведения', level=2)
    info_fields = {
        "ФИО": patient.get("ФИО", "—"),
        "Возраст": f"{patient.get('Возраст', '—')} лет",
        "Пол": patient.get("Пол", "—"),
        "Рост": f"{patient.get('Рост', '—')} см",
        "Вес": f"{patient.get('Вес', '—')} кг",
        "Сатурация (SpO₂)": f"{patient.get('spo2', '—')} %"
    }
    for k, v in info_fields.items():
        document.add_paragraph(f"{k}: {v}")

    # Шкалы
    document.add_heading('📊 Оценочные шкалы', level=2)

    def add_scale(doc, name, value):
        if isinstance(value, tuple):
            score, risk = value
            doc.add_paragraph(f"{name}: {score} — риск: {risk}")
        elif isinstance(value, bool):
            doc.add_paragraph(f"{name}: {'Высокий риск' if value else 'Низкий риск'}")
        else:
            doc.add_paragraph(f"{name}: {value}")

    add_scale(document, "Caprini", scales.get("caprini"))
    add_scale(document, "El-Ganzouri", scales.get("elganzouri"))
    add_scale(document, "Lee", scales.get("lee"))
    add_scale(document, "ARISCAT", scales.get("ariscat"))
    add_scale(document, "Индекс массы тела (BMI)", round(
        scales.get("bmi", [24.9, "Нормальная масса тела"])[0], 2))
    add_scale(document, "STOP-BANG", scales.get("stopbang"))
    add_scale(document, "SOBA", scales.get("soba"))

    # Сохранение в BytesIO
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    st.download_button(
        label="📄 Скачать DOCX-отчёт",
        data=file_stream,
        file_name=f"Отчет_{patient.get('ФИО', 'пациент')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    show_button_down()
